#!/usr/bin/env python3
"""Собирает Word-отчет из Markdown-аудита Bitrix24."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/pro2kuror/Desktop/architect")
SOURCE_MD = ROOT / "docs" / "architecture" / "bitrix24_commercial_director_audit_2026-05-01.md"
OUTPUT_DOCX = ROOT / "output" / "doc" / "bitrix24_commercial_director_audit_2026-05-01.docx"


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("<br>", "\n")
    return text.strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for index, part in enumerate(str(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip()
        cells = [clean_inline(cell) for cell in raw.strip("|").split("|")]
        if not all(re.fullmatch(r"\s*:?-{3,}:?\s*", cell or "") for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Отчетная таблица"
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for col_index in range(width):
            cell = row.cells[col_index]
            value = row_data[col_index] if col_index < len(row_data) else ""
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, "E7EEF8")
    document.add_paragraph()


def ensure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.color.rgb = RGBColor(31, 57, 92)
        style.font.bold = True

    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)

    if "Отчетная таблица" not in styles:
        table_style = styles.add_style("Отчетная таблица", WD_STYLE_TYPE.TABLE)
        table_style.base_style = styles["Table Grid"]


def set_page(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)


def add_paragraph(document: Document, text: str) -> None:
    text = clean_inline(text)
    if not text:
        document.add_paragraph()
        return
    if text.startswith("- "):
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(text[2:])
        return
    paragraph = document.add_paragraph()
    paragraph.add_run(text)


def build_docx() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    document = Document()
    set_page(document)
    ensure_styles(document)

    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue
        if stripped.startswith("# "):
            document.add_heading(clean_inline(stripped[2:]), level=0)
        elif stripped.startswith("## "):
            document.add_heading(clean_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            document.add_heading(clean_inline(stripped[4:]), level=2)
        else:
            add_paragraph(document, stripped)
        index += 1

    footer = document.sections[0].footer.paragraphs[0]
    footer.text = "Аудит Bitrix24. Подготовлено в МСК. Документ сформирован без изменений в Bitrix24."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(90, 90, 90)

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
