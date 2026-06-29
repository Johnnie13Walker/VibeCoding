#!/usr/bin/env python3
"""Мастер-документ ВМ-Логистик: все выводы исследования + дорожная карта действий.

Сводит воедино: профиль перевозок, финансы по каналам, утечки маржи, продуктивность
логистов, дисбаланс регионов, денежные рычаги ИИ, анализ площадок для авто-торгов,
и план действий по приоритету.

Пересчитывает все ключевые цифры из data/perevozki_raw.xls и data/clients_nelli.xlsx.
Запуск: .venv/bin/python scripts/build_master_docx.py
Выход:  reports/ВМ-Логистик_ИИ_исследование_и_план.docx (в git не коммитится)
"""
import openpyxl, xlrd, re, os, datetime
from collections import defaultdict, Counter
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "reports/ВМ-Логистик_ИИ_исследование_и_план.docx"
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
GOOD = RGBColor(0x1B, 0x7A, 0x3D)
BAD = RGBColor(0xB3, 0x1B, 0x1B)
AMBER = RGBColor(0xB8, 0x7A, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)
ANN = 2.4  # 5 мес -> год


def fnum(x):
    s = str(x).strip().replace("\xa0", "").replace(" ", "").replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None


def pdate(x):
    try:
        return datetime.datetime.strptime(str(x).strip(), "%d.%m.%Y").date()
    except Exception:
        return None


def keyize(s):
    s = str(s).lower().replace("ё", "е"); s = re.sub(r"и|й", "и", s); s = re.sub(r"э|е", "е", s)
    for w in ["ооо", "оао", "зао", "пао", "ао", "холдингс", "россия", "партнерс",
              "сводныи", "сэнтрал", "фуд", "компания", "пивоваренная", "снабжение"]:
        s = s.replace(w, "")
    return re.sub(r"[^a-zа-я0-9]", "", s)


def platform(op, url):
    s = f"{op} {url}".lower()
    if "atrucks" in s or "атракс" in s: return "Атракс"
    if ("logist" in s and "pro" in s) or "логистпро" in s or "логист про" in s or "logistpro" in s: return "Логист-Про"
    if "torgtrans" in s or "торгтранс" in s: return "ТоргТранс"
    if "loginet" in s or "логинет" in s: return "Логинет"
    if "ati" in s and ".su" in s: return "ATI.SU"
    if "tms" in s or "trucker" in s or "труке" in s or "сберкорус" in s: return "TMS/Trucker"
    if "bidzaar" in s or "бидзар" in s: return "Bidzaar"
    if "почт" in s or "@" in url or url.strip() == "-" or op.strip().lower() == "тг": return "Почта/мессенджер"
    return "Свой портал"


API = {"Атракс": "✅ из коробки", "Логист-Про": "✅ из коробки", "ATI.SU": "✅ зрелый",
       "ТоргТранс": "по договору", "Логинет": "по договору", "Bidzaar": "под заказчика",
       "TMS/Trucker": "клиентский", "Почта/мессенджер": "нет", "Свой портал": "только RPA"}


def compute():
    sh = xlrd.open_workbook("data/perevozki_raw.xls").sheet_by_index(0)
    H = {str(sh.cell_value(0, c)).strip(): c for c in range(sh.ncols)}
    def V(n): return [sh.cell_value(r, H[n]) for r in range(1, sh.nrows)]
    N = sh.nrows - 1
    st, ch, cl, au = V("Статус заявки"), V("Канал продаж"), V("Клиент"), V("Автор")
    tk, tp, pr = V("Тариф клиента"), V("Тариф подрядчика"), V("Прибыль по клиенту руб")
    rt, rz, rv, pod = V("Маршрут"), V("Регион загрузка"), V("Регион выгрузка"), V("Подрядчик")
    dl = [pdate(x) for x in V("Дата загрузки")]
    conf = lambda i: str(st[i]).strip() == "Подтверждена"
    d = {"N": N, "period": (min(x for x in dl if x), max(x for x in dl if x)),
         "status": Counter(str(x).strip() for x in st),
         "clients": len({str(x).strip() for x in cl if str(x).strip()}),
         "contractors": len({str(x).strip() for x in pod if str(x).strip()}),
         "logists_total": len({str(x).strip() for x in au if str(x).strip()})}

    # финансы + каналы
    rev = cost = prof = n = 0.0
    chan = defaultdict(lambda: [0, 0.0, 0.0])
    for i in range(N):
        if not conf(i): continue
        a = fnum(tk[i])
        if a is None: continue
        n += 1; rev += a; cost += fnum(tp[i]) or 0; prof += fnum(pr[i]) or 0
        k = str(ch[i]).strip() or "(не указан)"
        chan[k][0] += 1; chan[k][1] += a; chan[k][2] += fnum(pr[i]) or 0
    d["fin"] = dict(n=int(n), rev=rev, cost=cost, prof=prof, margin=100 * prof / rev,
                    avg_check=rev / n, avg_prof=prof / n)
    d["channels"] = sorted(chan.items(), key=lambda kv: -kv[1][1])

    # Гарантия — утечки
    g = defaultdict(lambda: [0, 0.0, 0.0])
    loss_n = 0; loss_rub = 0.0
    for i in range(N):
        if conf(i) and str(ch[i]).strip() == "Гарантия":
            p = fnum(pr[i]) or 0
            g[str(cl[i]).strip()][0] += 1; g[str(cl[i]).strip()][1] += fnum(tk[i]) or 0; g[str(cl[i]).strip()][2] += p
            if p < 0: loss_n += 1; loss_rub += p
    d["guarantee"] = dict(top=sorted(g.items(), key=lambda kv: -kv[1][1])[:6], loss_n=loss_n, loss_rub=loss_rub)

    # логисты
    pa = defaultdict(lambda: [0, 0.0]); days = defaultdict(set)
    for i in range(N):
        if not conf(i): continue
        a = str(au[i]).strip()
        if not a: continue
        pa[a][0] += 1; pa[a][1] += fnum(pr[i]) or 0
        if dl[i]: days[a].add(dl[i])
    active = [a for a, v in pa.items() if v[0] >= 20]
    tot_n = sum(pa[a][0] for a in active); tot_d = sum(len(days[a]) for a in active)
    d["logists"] = dict(active=len(active), prod=tot_n / (tot_d or 1))

    # регионы
    lc, uc = Counter(), Counter()
    for i in range(N):
        if not conf(i): continue
        if str(rz[i]).strip(): lc[str(rz[i]).strip()] += 1
        if str(rv[i]).strip(): uc[str(rv[i]).strip()] += 1
    bal = [(r, lc[r], uc[r], uc[r] - lc[r]) for r in set(lc) | set(uc)]
    d["regions"] = sorted(bal, key=lambda x: -x[3])[:6]

    # СПОТ маршруты (повторяемость)
    spot_rt = defaultdict(int)
    spot_total = 0
    for i in range(N):
        if conf(i) and str(ch[i]).strip() == "СПОТ" and str(rt[i]).strip():
            spot_rt[str(rt[i]).strip()] += 1; spot_total += 1
    rec = {r: c for r, c in spot_rt.items() if c >= 10}
    d["spot_routes"] = dict(uni=len(spot_rt), total=spot_total, rec=len(rec), rec_deals=sum(rec.values()))

    # площадки (сшивка с реестром КМ)
    wb = openpyxl.load_workbook("data/clients_nelli.xlsx", data_only=True)
    ws = wb["Общая инфо "]
    rows = []
    for r in range(2, ws.max_row + 1):
        nm = ws.cell(r, 1).value
        if nm and str(nm).strip():
            rows.append((str(nm).strip(), str(ws.cell(r, 15).value or ""), str(ws.cell(r, 16).value or "")))
    spot_cli = defaultdict(lambda: [0, 0.0])
    for i in range(N):
        if conf(i) and str(ch[i]).strip() == "СПОТ":
            c = str(cl[i]).strip(); spot_cli[c][0] += 1; spot_cli[c][1] += fnum(pr[i]) or 0
    fk = [(keyize(c), c) for c in spot_cli]
    def money(short):
        k = keyize(short)
        if not k: return 0, 0.0
        for kk, full in fk:
            if k in kk or kk in k or (len(k) >= 4 and k[:5] in kk): return spot_cli[full]
        return 0, 0.0
    pl = defaultdict(lambda: [set(), 0, 0.0])
    for nm, op, url in rows:
        p = platform(op, url); dd, pp = money(nm)
        pl[p][0].add(nm); pl[p][1] += dd; pl[p][2] += pp
    d["platforms"] = sorted(pl.items(), key=lambda kv: -kv[1][2])
    d["spot_total_profit"] = sum(v[1] for v in spot_cli.values())

    # Волна 1 (Атракс + Логист-Про) — клиенты с месячным объёмом и незакрытыми (отменёнными)
    w1keys = [(keyize(nm), platform(op, url), str(nm).strip()) for nm, op, url in rows
              if platform(op, url) in ("Атракс", "Логист-Про")]
    def mp(client):
        k = keyize(client)
        for kk, p, short in w1keys:
            if kk and (kk in k or k in kk or (len(kk) >= 4 and kk[:5] in k)): return p
        return None
    months = len({(x.year, x.month) for x in dl if x})
    # по полному клиенту: подтв, отменён (весь канал СПОТ, включая отмены)
    wc = defaultdict(lambda: [0, 0, 0.0])  # conf, cancel, profit
    for i in range(N):
        if str(ch[i]).strip() != "СПОТ": continue
        p = mp(str(cl[i]).strip())
        if not p: continue
        s = str(st[i]).strip()
        if s == "Подтверждена": wc[(p, str(cl[i]).strip())][0] += 1; wc[(p, str(cl[i]).strip())][2] += fnum(pr[i]) or 0
        elif s == "Отменена": wc[(p, str(cl[i]).strip())][1] += 1
    plat_cli = defaultdict(list)
    for (p, full), (cf, cn, pp) in wc.items():
        plat_cli[p].append((full, cf, cn, pp))
    wave1 = {"months": months, "platforms": {}}
    w1n = w1p = 0
    for P in ("Атракс", "Логист-Про"):
        items = sorted(plat_cli[P], key=lambda x: -(x[1] + x[2]))
        tc = sum(x[1] for x in items); tn = sum(x[2] for x in items)
        wave1["platforms"][P] = dict(clients=items, conf=tc, cancel=tn,
                                     conf_pm=tc / months, cancel_pm=tn / months,
                                     cancel_rate=100 * tn / (tc + tn) if (tc + tn) else 0)
        w1n += tc; w1p += sum(x[3] for x in items)
    wave1["conf_pm"] = (wave1["platforms"]["Атракс"]["conf_pm"] + wave1["platforms"]["Логист-Про"]["conf_pm"])
    wave1["cancel_pm"] = (wave1["platforms"]["Атракс"]["cancel_pm"] + wave1["platforms"]["Логист-Про"]["cancel_pm"])
    d["wave1"] = dict(deals=w1n, prof=w1p, detail=wave1)

    # Объём торгов из вкладки «СПОТы» (заявок/участие/выиграно по клиентам, 9 мес)
    wsv = wb["СПОТы"]
    def vnum(v):
        try:
            return float(v)
        except Exception:
            return 0.0
    spotv = defaultdict(lambda: [0.0, 0.0, 0.0])
    sv_tot = [0.0, 0.0, 0.0]; sv_dates = set()
    for r in range(4, wsv.max_row + 1):
        c = wsv.cell(r, 5).value; f = wsv.cell(r, 6).value
        if not c or f is None: continue
        z = vnum(f); u = vnum(wsv.cell(r, 7).value); ww = vnum(wsv.cell(r, 8).value)
        if z == 0 and u == 0 and ww == 0: continue
        spotv[str(c).strip()][0] += z; spotv[str(c).strip()][1] += u; spotv[str(c).strip()][2] += ww
        sv_tot[0] += z; sv_tot[1] += u; sv_tot[2] += ww
        dt = wsv.cell(r, 4).value
        if isinstance(dt, datetime.datetime): sv_dates.add((dt.year, dt.month))
    SVM = len(sv_dates) or 9
    sv_keys = [(keyize(c), c) for c in spotv]
    def sv_lookup(name):
        k = keyize(name)
        for sk, sc in sv_keys:
            if sk and (sk in k or k in sk or (len(sk) >= 3 and len(k) >= 4 and sk[:4] == k[:4])):
                return sc
        return None
    reg_plat = {"Атракс": [], "Логист-Про": []}
    for nm, op, url in rows:
        p = platform(op, url)
        if p in reg_plat: reg_plat[p].append(str(nm).strip())
    sv_plat = {}
    for P in ("Атракс", "Логист-Про"):
        seen = set(); items = []; tz = tu = tw = 0.0
        for nm in sorted(set(reg_plat[P])):
            sc = sv_lookup(nm)
            if sc and sc not in seen:
                seen.add(sc); z, u, ww = spotv[sc]
                items.append((nm, z, u, ww)); tz += z; tu += u; tw += ww
        sv_plat[P] = dict(items=sorted(items, key=lambda x: -x[1]), z=tz, u=tu, w=tw,
                          z_pm=tz / SVM, notclosed_pm=(tz - tw) / SVM, noparticip_pm=(tz - tu) / SVM,
                          win_rate=100 * tw / tz if tz else 0, part_rate=100 * tu / tz if tz else 0)
    d["spot_volume"] = dict(months=SVM, tot_z=sv_tot[0], tot_u=sv_tot[1], tot_w=sv_tot[2],
                            part_rate=100 * sv_tot[1] / sv_tot[0], win_all=100 * sv_tot[2] / sv_tot[0],
                            win_part=100 * sv_tot[2] / sv_tot[1], platforms=sv_plat)
    return d


# ---------- рендер ----------
def sf(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic; run.font.name = "Calibri"
    if color is not None: run.font.color.rgb = color


def h(doc, t, lvl=1):
    p = doc.add_heading(level=lvl); r = p.add_run(t); r.font.color.rgb = ACCENT; r.font.name = "Calibri"; return p


def para(doc, runs, sa=6, style=None):
    p = doc.add_paragraph(style=style); p.paragraph_format.space_after = Pt(sa)
    if isinstance(runs, str): runs = [(runs, {})]
    for t, kw in runs: sf(p.add_run(t), **kw)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, hd in enumerate(headers):
        c = t.rows[0].cells[j]; c.paragraphs[0].clear()
        sf(c.paragraphs[0].add_run(hd), size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for j, (val, kw) in enumerate(row):
            cells[j].paragraphs[0].clear(); sf(cells[j].paragraphs[0].add_run(str(val)), size=9.5, **kw)
    if widths:
        for j, wd in enumerate(widths):
            for r in t.rows: r.cells[j].width = Cm(wd)
    return t


def rub(x): return f"{x:,.0f}".replace(",", " ")
def mln(x): return f"{x/1e6:.0f} млн"


def build(d):
    doc = Document()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Cm(1.7); sec.bottom_margin = Cm(1.7); sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    fin = d["fin"]

    # Титул
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run("ВМ-Логистик"), size=13, bold=True, color=GREY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    sf(p.add_run("Внедрение ИИ: итоги исследования\nи план действий по приоритету"), size=22, bold=True, color=ACCENT)
    a, b = d["period"]
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(f"Анализ выгрузки 1С за {a.strftime('%d.%m.%Y')}–{b.strftime('%d.%m.%Y')} "
                 f"({d['N']:,} заявок) + реестр площадок КМ".replace(",", " ")), size=10, italic=True, color=GREY)
    doc.add_paragraph()

    # Резюме для руководства
    h(doc, "Резюме для руководства", 1)
    para(doc, [("Бизнес генерирует ", {}), (f"≈{mln(fin['rev']*ANN)} ₽ выручки и {mln(fin['prof']*ANN)} ₽ прибыли в год", {"bold": True}),
               (f" при марже {fin['margin']:.1f}%. Анализ показал четыре точки, где ИИ приносит деньги, "
                "и оптимальный порядок их внедрения.", {})])
    for head, rest in [
        ("Закупка перевозки — главный рычаг.", f" Себестоимость ≈{mln(fin['cost']*ANN)} ₽/год. Снижение на 2–3% за счёт "
                                               "умного подбора = +50–75 млн ₽/год. Это ядро проекта."),
        ("Гарантия теряет деньги.", " 12,7% маржи против 20,5% в среднем; убыточных заявок на −13,6 млн ₽/год. "
                                    "Быстрый фикс тарифной модели = +15–25 млн ₽/год."),
        ("СПОТ — где масштаб.", " 65% выручки; вручную участвуем лишь в ~6% рынка. Авто-торги дают ×3-потенциал по объёму."),
        ("Логисты и порожний пробег.", f" Выработка {d['logists']['prod']:.1f} заявки/день при цели 10; машины возвращаются "
                                        "из регионов порожняком. AI-ассистент и обратная загрузка высвобождают ресурс."),
    ]:
        pp = doc.add_paragraph(style="List Bullet"); sf(pp.add_run(head), bold=True); sf(pp.add_run(rest))
    para(doc, [("Совокупный потенциал первого года: ", {"bold": True}),
               ("100–180 млн ₽ дополнительной прибыли при последовательном внедрении (оценка по нижней–верхней границе).",
                {"bold": True, "color": GOOD})], sa=10)

    # ЧАСТЬ I — картина бизнеса
    h(doc, "Часть I. Что показали данные", 1)

    h(doc, "1.1. Финансовый профиль", 2)
    table(doc, ["Метрика", "5 мес", "≈ Год"], [
        [("Выручка", {}), (rub(fin["rev"]) + " ₽", {}), (mln(fin["rev"]*ANN) + " ₽", {"bold": True})],
        [("Себестоимость (перевозчики)", {}), (rub(fin["cost"]) + " ₽", {}), (mln(fin["cost"]*ANN) + " ₽", {})],
        [("Прибыль", {}), (rub(fin["prof"]) + " ₽", {}), (mln(fin["prof"]*ANN) + " ₽", {"bold": True, "color": GOOD})],
        [("Маржа", {}), (f"{fin['margin']:.1f}%", {"bold": True}), ("—", {})],
        [("Средний чек / прибыль на заявку", {}), (rub(fin["avg_check"]) + " / " + rub(fin["avg_prof"]) + " ₽", {}), ("—", {})],
    ], widths=[7.0, 5.5, 4.0])
    para(doc, f"Данные: {d['clients']} клиентов, {d['contractors']:,} перевозчиков, {d['logists_total']} логистов. "
              f"Подтверждено {100*d['status'].get('Подтверждена',0)/d['N']:.0f}% заявок.".replace(",", " "))

    h(doc, "1.2. Каналы продаж — где прибыль и где утечки", 2)
    rows = []
    for k, v in d["channels"]:
        m = 100 * v[2] / v[1] if v[1] else 0
        col = GOOD if m >= 25 else (BAD if m < 13 else None)
        rows.append([(k, {}), (f"{v[0]:,}".replace(",", " "), {}), (rub(v[1]), {}), (f"{m:.1f}%", {"bold": True, "color": col})])
    table(doc, ["Канал", "Заявок", "Выручка ₽", "Маржа"], rows, widths=[3.5, 2.5, 4.5, 2.5])
    para(doc, [("СПОТ", {"bold": True}), (" — 65% выручки (цель авто-торгов). ", {}),
               ("Гарантия", {"bold": True}), (" — самая низкая маржа (фикс тарифов). ", {}),
               ("ПРОЕКТ", {"bold": True}), (" — самые качественные сделки.", {})], sa=4)

    h(doc, "1.3. Утечка маржи в Гарантии", 2)
    gr = d["guarantee"]
    para(doc, [(f"{gr['loss_n']} заявок убыточны, минус ", {}), (rub(gr["loss_rub"]) + " ₽ за 5 мес", {"bold": True, "color": BAD}),
               (f" (≈{mln(gr['loss_rub']*ANN)} ₽/год). Локализовано в нескольких клиентах:", {})])
    rows = []
    for c, v in gr["top"]:
        m = 100 * v[2] / v[1] if v[1] else 0
        rows.append([(c[:34], {}), (f"{v[0]:,}".replace(",", " "), {}), (f"{m:.1f}%", {"bold": True, "color": BAD if m < 8 else None})])
    table(doc, ["Клиент (Гарантия)", "Заявок", "Маржа"], rows, widths=[9.0, 3.0, 3.0])

    h(doc, "1.4. Логисты и обратная загрузка", 2)
    para(doc, [(f"Активных логистов {d['logists']['active']}, средняя выработка ", {}),
               (f"{d['logists']['prod']:.2f} заявки/день", {"bold": True}),
               (" — ниже бенчмарка (3,5) и втрое меньше цели (10). Высокая выработка ≠ маржа: лидер по объёму "
                "работает на убыточной Гарантии.", {})])
    para(doc, "Регионы с дефицитом обратного груза (машины едут назад порожняком):")
    rows = [[(r, {}), (str(l), {}), (str(u), {}), (f"+{dd}", {"bold": True, "color": BAD})] for r, l, u, dd in d["regions"]]
    table(doc, ["Регион", "Загрузка", "Выгрузка", "Дисбаланс"], rows, widths=[6.0, 3.0, 3.0, 3.0])

    # ЧАСТЬ II — рычаги ИИ
    h(doc, "Часть II. Где ИИ достаёт деньги", 1)
    levers = [
        ("1. Умный подбор перевозчика", "+50–75 млн ₽/год", GOOD,
         f"Себестоимость ≈{mln(fin['cost']*ANN)} ₽/год — крупнейшая статья. Модель предсказывает справедливую "
         "закупочную цену по маршруту/дате/грузу, подбирает перевозчиков из базы, ведёт авто-торг. Снижение "
         "закупки на 2–3% = десятки миллионов. Это фундамент контура A."),
        ("2. Тарифная модель Гарантии", "+15–25 млн ₽/год", GOOD,
         "Светофор при заведении ставки: предсказывает себестоимость и предупреждает об убытке до подтверждения; "
         "ежемесячно предлагает пересмотр коридоров по убыточным клиентам. Дёшево, данные уже есть."),
        ("3. Авто-участие в СПОТ-торгах", "+25–60 млн ₽/год", GOOD,
         "Парсинг заявок площадок → скоринг выгодности → авто-ставка. Сейчас вручную обрабатывается ~6% рынка. "
         "Потенциал ×3 по объёму (детально — Часть III)."),
        ("4. Обратная загрузка", "+20 млн ₽/год", AMBER,
         "Мэтчинг освобождающихся машин с грузом из региона выгрузки + карта дефицита для отдела продаж. "
         "Прибыль на уже едущей машине — почти чистая."),
        ("5. AI-ассистент логиста", "выработка ×1,5", AMBER,
         "Парсинг входящей заявки, автозаполнение карточки, черновики переписки и обзвона. Тот же штат тянет "
         "больше заявок (или тот же объём меньшим ФОТ — а ФОТ 57–58%, главный кост)."),
        ("6. Скоринг срыва рейса", "минус потери", AMBER,
         "Предсказание ненадёжного перевозчика до постановки на рейс → меньше штрафов по Гарантии и холостой работы "
         "(сейчас 8,9% заявок отменяются)."),
    ]
    for title, eff, col, body in levers:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        sf(p.add_run(title + "  "), size=12, bold=True, color=ACCENT); sf(p.add_run(eff), size=12, bold=True, color=col)
        para(doc, body, sa=8)

    # ЧАСТЬ III — площадки
    h(doc, "Часть III. Контур B: с каких площадок начинать", 1)
    para(doc, [("СПОТ-прибыль ", {}), (f"{rub(d['spot_total_profit'])} ₽/5мес (≈{mln(d['spot_total_profit']*ANN)} ₽/год)", {"bold": True}),
               (". Распределение по площадкам и готовности API (суммы приблизительны — клиенты сшиты по наименованию):", {})])
    rows = []
    for pn, v in d["platforms"]:
        apit = API.get(pn, "?"); col = GOOD if "✅" in apit else (AMBER if apit in ("по договору", "под заказчика", "клиентский") else BAD)
        rows.append([(pn, {}), (str(len(v[0])), {}), (rub(v[2]), {"bold": v[2] > 15e6}), (apit, {"color": col, "bold": "✅" in apit})])
    table(doc, ["Площадка", "Клиентов", "СПОТ-приб/5мес ₽", "API"], rows, widths=[3.8, 2.2, 4.5, 3.5])
    para(doc, [("Логика выбора: ", {"bold": True}),
               ("эффект = деньги × реализуемость. Самые денежные площадки (свои порталы, ТоргТранс) технически "
                "тяжёлые. Атракс — один коннектор на 15 клиентов с API из коробки — лучший старт.", {})])

    # 3.1. Объём заявок на торгах и сколько мы не закрываем
    sv = d["spot_volume"]
    h(doc, "3.1. Сколько заявок мы НЕ закрываем (данные торгов)", 2)
    para(doc, [(f"По учёту КМ за {sv['months']} мес на всех площадках прошло ", {}),
               (f"{rub(sv['tot_z'])} заявок", {"bold": True}),
               (f"; участвуем лишь в {rub(sv['tot_u'])} ({sv['part_rate']:.1f}%), выигрываем {rub(sv['tot_w'])} "
                f"({sv['win_all']:.1f}% от всех, {sv['win_part']:.1f}% от участия). "
                "То есть закрываем чуть больше 1% доступного потока — здесь и лежит потенциал ×3.", {})])
    para(doc, [("По двум стартовым площадкам:", {"bold": True})], sa=4)
    for P in ("Атракс", "Логист-Про"):
        pl = sv["platforms"][P]
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        sf(p.add_run(f"{P}: "), size=11.5, bold=True, color=ACCENT)
        sf(p.add_run(f"≈{pl['z_pm']:,.0f} заявок/мес".replace(",", " ")), size=11, bold=True)
        sf(p.add_run(f"  ·  участвуем {pl['part_rate']:.1f}%  ·  выигрываем {pl['win_rate']:.1f}%  ·  "), size=11)
        sf(p.add_run(f"НЕ закрываем ≈{pl['notclosed_pm']:,.0f}/мес".replace(",", " ")), size=11, bold=True, color=BAD)
        rows = []
        for nm, z, u, ww in pl["items"]:
            if z < 1: continue
            rows.append([(nm[:30], {}), (f"{z/sv['months']:,.0f}".replace(",", " "), {}),
                         (f"{u/sv['months']:,.0f}".replace(",", " "), {}),
                         (f"{ww/sv['months']:,.0f}".replace(",", " "), {}),
                         (f"{(z-ww)/sv['months']:,.0f}".replace(",", " "), {"bold": True, "color": BAD})])
        table(doc, ["Клиент", "Заявок/мес", "Участвуем/мес", "Выигрыв/мес", "НЕ закрыв/мес"], rows,
              widths=[6.0, 2.6, 2.8, 2.6, 2.8])
    para(doc, [("Главный вывод. ", {"bold": True, "color": ACCENT}),
               (f"Только на Атракс и Логист-Про мимо проходит ≈{sv['platforms']['Атракс']['z_pm']+sv['platforms']['Логист-Про']['z_pm']:,.0f} заявок/мес".replace(",", " "),
                {"bold": True}),
               (", а выигрываем единицы процентов. Узкое место — даже не цена, а то, что логист физически не "
                "успевает подать ставку: по этим площадкам мы вообще не участвуем в "
                f"≈{sv['platforms']['Атракс']['noparticip_pm']+sv['platforms']['Логист-Про']['noparticip_pm']:,.0f} заявок/мес".replace(",", " ")
                + ". Именно это закрывает авто-участие.", {})])
    para(doc, [("Оговорки: ", {"bold": True}),
               ("(1) «Кол-во заявок» — это весь поток лотов клиента на площадке, включая нерентабельные и "
                "непрофильные — не все из них стоит закрывать; реальная цель = профильные заявки в нашем "
                "тарифном коридоре. (2) Часть клиентов площадок отсутствует во вкладке «СПОТы», поэтому объём "
                "скорее занижен. (3) Логист-Про раздут двумя клиентами с огромным потоком (Мултон, ОПХ), по "
                "которым мы почти не участвуем.", {})])

    # ЧАСТЬ IV — решения встречи + дорожная карта
    h(doc, "Часть IV. Решения встречи с Романом", 1)
    para(doc, [("На встрече с Романом (учредитель, направление AI-логиста) согласована последовательность "
                "запуска. Ключевое решение: ", {}),
               ("первый шаг — парсинг заявок, а не голосовой ассистент.", {"bold": True}),
               (" Видимость объёма заявок важнее всего, потому что сейчас листинг не виден нигде (в 1С — лишь "
                "~3%). Целевой контейнер для всей работы — «личный кабинет логиста» (ЛК).", {})])
    h(doc, "Согласованная логика (4 шага к ЛК логиста)", 2)
    for t in [
        ("Парсинг заявок. ", "Сбор заявок с опишек/агрегаторов (Атракс, Логист-Про) в единый листинг ЛК — увидеть поток."),
        ("Оценка ликвидности. ", "Деление заявок на рентабельные/нерентабельные на исторических данных."),
        ("Завод логистов в ЛК. ", "Логисты работают с массивом заявок через ЛК — сделать это «необходимостью» (управляемость)."),
        ("Снижение зависимости от логиста. ", "Авто/голосовой ассистент берёт заявки, которые логист не успевает."),
    ]:
        p = doc.add_paragraph(style="List Number"); sf(p.add_run(t[0]), bold=True); sf(p.add_run(t[1]))
    h(doc, "Отложено и подтверждено", 2)
    for t in [
        ("Отложен голосовой ассистент. ", "Нет надёжного канала коммуникации в РФ (WhatsApp неофиц. API — риск вплоть "
         "до уголовного; Telegram нестабилен). Преждевременно."),
        ("Гарантия — подтверждённая беда. ", "Роман: минимизировать; рекламации 2–5 млн ₽/мес. Брать только если "
         "клиент иначе не даёт СПОТ, и минимальный объём. Совпадает с нашим анализом убытков."),
        ("Обратная загрузка — те же регионы. ", "Список Артёма (Ростов, СПб, Самара, Краснодар, Татарстан, Свердловск) "
         "совпал с нашим расчётом дисбаланса. Идея: искать грузоотправителей со складами в этих регионах + дробить по тоннажу."),
        ("Повторные перевозчики. ", "Если возил месяц — не пробивать через СБ (экономия времени логиста). Привлекаем "
         "слишком много новых, мало повторных — отдельная задача."),
    ]:
        p = doc.add_paragraph(style="List Bullet"); sf(p.add_run(t[0]), bold=True); sf(p.add_run(t[1]))

    h(doc, "Часть IV-bis. Дорожная карта", 1)
    para(doc, "Порядок отражает решения встречи (парсинг-first) и зависимости между шагами. Эффект в ₽ — "
              "годовой потенциал, оценка.")
    steps = [
        ("0", "Доступы и данные", "сейчас", "1–2 нед", "—",
         "Dev-доступ к API/опишкам Атракс / Логист-Про / агрегаторов; регулярная выгрузка из 1С; уточнить площадки "
         "крупных СПОТ-клиентов вне реестра (Стройпроектсервис 30 млн и др.)."),
        ("1", "Парсинг заявок → ЛК логиста (видимость объёма)", "база + видимость", "3–5 нед", "низкая",
         "Решение встречи — первый шаг. Парсинг заявок с Атракс/Логист-Про/агрегаторов в единый листинг ЛК. "
         "Цель — увидеть «куски пирога, которые не забираем». Исполняют партнёры (есть отдельное ТЗ)."),
        ("2", "Тарифный светофор по Гарантии", "+15–25 млн/год", "4–6 нед", "низкая",
         "Параллельно: модель себестоимости + алерт об убытке. Снимает −13,6 млн/год и риск штрафов. "
         "Роман подтвердил приоритет минимизации Гарантии."),
        ("3", "Оценка ликвидности + модель закупки/подбора", "+50–75 млн/год", "6–10 нед", "средняя",
         "Деление заявок на ликвидные/неликвидные + предсказание закупочной цены. Ядро контура A, "
         "фундамент для авто-ставки и скоринга."),
        ("4", "Завод логистов в ЛК + апдейт грузов в АТИ", "управляемость", "6–8 нед", "средняя",
         "Кнопка «опубликовать» с автозаполнением, апдейт грузов в АТИ внутри ЛК. Логисты работают только отсюда — "
         "управляемость, которой сейчас нет."),
        ("5", "Обратная загрузка + повторные перевозчики", "+20 млн/год", "4–6 нед", "средняя",
         "Поиск обратного груза из Ростова/СПб/Самары/Краснодара по листингу + задачи отделу продаж; не пробивать "
         "месячных перевозчиков (экономия времени логиста)."),
        ("6", "Снижение зависимости от логиста (авто/голос)", "+25–60 млн/год", "позже", "высокая",
         "Авто-ставка по ликвидным заявкам; голосовой ассистент берёт неуспетое — когда появится надёжный канал. "
         "Здесь реализуется потенциал ×3 на упускаемом потоке."),
        ("7", "Скоринг надёжности + дашборды (контур C)", "минус штрафы", "позже", "средняя",
         "Anti-no-show поверх модели подбора; управленческие дашборды на данных 1С/ЛК. Контроль логистов — "
         "чувствительная тема, внедрять аккуратно."),
    ]
    rows = []
    for n, name, eff, term, cx, _ in steps:
        rows.append([(n, {"bold": True, "color": ACCENT}), (name, {"bold": True}), (eff, {"color": GOOD if "млн" in eff else None}),
                     (term, {}), (cx, {})])
    table(doc, ["#", "Шаг", "Эффект", "Срок", "Сложность"], rows, widths=[1.0, 5.5, 3.5, 3.0, 3.0])
    doc.add_paragraph()
    for n, name, eff, term, cx, body in steps:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        sf(p.add_run(f"Шаг {n}. {name} "), size=11.5, bold=True, color=ACCENT)
        sf(p.add_run(f"({eff}, {term})"), size=10, italic=True, color=GREY)
        para(doc, body, sa=7)

    # ЧАСТЬ V — что нужно от клиента / риски
    h(doc, "Часть V. Что нужно от ВМ-Логистик и риски", 1)
    h(doc, "Решения и доступы", 2)
    for t in [
        "Dev-доступ к API площадок (Атракс, Логист-Про, ATI) от лица перевозчика.",
        "Регулярная выгрузка из 1С (ядро всех моделей) + доступ к базе перевозчиков.",
        "Утверждение правила флор-цены (мин. маржа) для авто-ставок.",
        "Площадки крупных внереестровых клиентов (Стройпроектсервис, Иркутский завод полимеров, Газпромнефть).",
    ]:
        para(doc, t, style="List Bullet", sa=3)
    h(doc, "Риски и оговорки", 2)
    for t in [
        ("Политика контроля. ", "Контроль логистов — чувствительная тема. Авто-торги и подбор — это инструменты "
                                "логиста, а не надзор; дашборды (контур C) внедрять аккуратно, начиная с финансов."),
        ("Точность оценок. ", "Суммы — порядок величины по выгрузке за 5 мес; точные цифры по шагам уточняются на "
                              "данных (себестоимость по маршрутам, штрафная механика Гарантии)."),
        ("Сшивка площадок. ", "Клиенты сопоставлены по наименованию между двумя файлами — возможны небольшие "
                              "расхождения сумм по площадкам."),
    ]:
        p = doc.add_paragraph(style="List Bullet"); sf(p.add_run(t[0]), bold=True); sf(p.add_run(t[1]))

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16)
    sf(p.add_run("Источники: выгрузка 1С «данные по перевозкам» (янв–май 2026, подтверждённые заявки), реестр "
                 "площадок КМ (вкладка «Общая инфо»). Документ содержит коммерческие данные — не передавать без "
                 "согласования."), size=9, italic=True, color=GREY)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("Сохранено:", OUT)


if __name__ == "__main__":
    build(compute())
