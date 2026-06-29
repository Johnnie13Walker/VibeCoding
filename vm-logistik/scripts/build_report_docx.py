#!/usr/bin/env python3
"""Сборка Word-отчёта по выгрузке 1С (ВМ-Логистик), янв–май 2026.

Пересчитывает метрики из data/perevozki_raw.xls и рендерит .docx с таблицами.
Запуск: .venv/bin/python scripts/build_report_docx.py
Выход:  reports/Анализ_перевозок_янв-май_2026.docx  (в git не коммитится)
"""
import xlrd, datetime, os
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PATH = "data/perevozki_raw.xls"
OUT = "reports/Анализ_перевозок_янв-май_2026.docx"

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)   # тёмно-синий
GOOD = RGBColor(0x1B, 0x7A, 0x3D)
BAD = RGBColor(0xB3, 0x1B, 0x1B)
GREY = RGBColor(0x60, 0x60, 0x60)


# ---------- данные ----------
def load():
    sh = xlrd.open_workbook(PATH).sheet_by_index(0)
    hdr = [str(sh.cell_value(0, c)).strip() for c in range(sh.ncols)]
    H = {h: i for i, h in enumerate(hdr)}
    N = sh.nrows - 1
    def V(name):
        c = H[name]
        return [sh.cell_value(r, c) for r in range(1, sh.nrows)]
    return V, N


def fnum(x):
    s = str(x).strip().replace("\xa0", "").replace(" ", "").replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None


def pdate(x):
    try:
        return datetime.datetime.strptime(str(x).strip(), "%d.%m.%Y").date()
    except Exception:
        return None


def rub(x):
    return f"{x:,.0f}".replace(",", " ")


def compute():
    V, N = load()
    st = V("Статус заявки"); ch = V("Канал продаж"); rt = V("Маршрут")
    cl = V("Клиент"); au = V("Автор"); tk = V("Тариф клиента"); tp = V("Тариф подрядчика")
    pr = V("Прибыль по клиенту руб"); rz = V("Регион загрузка"); rv = V("Регион выгрузка")
    pod = V("Подрядчик"); org = V("Организация"); dl = [pdate(x) for x in V("Дата загрузки")]
    conf = lambda i: str(st[i]).strip() == "Подтверждена"
    d = {}
    d["N"] = N
    d["period"] = (min(x for x in dl if x), max(x for x in dl if x))
    d["status"] = Counter(str(x).strip() for x in st)
    d["clients_uni"] = len({str(x).strip() for x in cl if str(x).strip()})
    d["contractors_uni"] = len({str(x).strip() for x in pod if str(x).strip()})
    d["authors_uni"] = len({str(x).strip() for x in au if str(x).strip()})
    d["org"] = Counter(str(x).strip() for x in org if str(x).strip())

    # итоги (подтв.)
    sk = sp = spr = n = 0.0
    for i in range(N):
        if not conf(i):
            continue
        a = fnum(tk[i])
        if a is None:
            continue
        n += 1; sk += a; sp += fnum(tp[i]) or 0; spr += fnum(pr[i]) or 0
    d["fin"] = dict(n=int(n), rev=sk, cost=sp, profit=spr, margin=100 * spr / sk,
                    avg_check=sk / n, avg_profit=spr / n)

    # каналы
    chan = defaultdict(lambda: [0, 0.0, 0.0])
    for i in range(N):
        if not conf(i):
            continue
        k = str(ch[i]).strip() or "(не указан)"
        chan[k][0] += 1; chan[k][1] += fnum(tk[i]) or 0; chan[k][2] += fnum(pr[i]) or 0
    d["channels"] = sorted(chan.items(), key=lambda kv: -kv[1][1])

    # помесячно
    mon = defaultdict(lambda: [0, 0.0, 0.0])
    for i in range(N):
        if not conf(i) or not dl[i]:
            continue
        k = (dl[i].year, dl[i].month)
        mon[k][0] += 1; mon[k][1] += fnum(tk[i]) or 0; mon[k][2] += fnum(pr[i]) or 0
    d["months"] = sorted(mon.items())

    # топ-клиенты по выручке
    cli = defaultdict(lambda: [0, 0.0, 0.0])
    for i in range(N):
        if not conf(i):
            continue
        k = str(cl[i]).strip()
        cli[k][0] += 1; cli[k][1] += fnum(tk[i]) or 0; cli[k][2] += fnum(pr[i]) or 0
    d["top_clients"] = sorted(cli.items(), key=lambda kv: -kv[1][1])[:12]

    # СПОТ-маршруты
    spot = defaultdict(lambda: [0, 0.0])
    for i in range(N):
        if conf(i) and str(ch[i]).strip() == "СПОТ" and str(rt[i]).strip():
            spot[str(rt[i]).strip()][0] += 1; spot[str(rt[i]).strip()][1] += fnum(pr[i]) or 0
    rec = {r: v for r, v in spot.items() if v[0] >= 10}
    d["spot"] = dict(uni=len(spot), total=sum(v[0] for v in spot.values()),
                     rec_n=len(rec), rec_deals=sum(v[0] for v in rec.values()),
                     rec_rev=sum(v[1] for v in rec.values()),
                     top=sorted(rec.items(), key=lambda kv: -kv[1][0])[:12])

    # Гарантия по клиентам
    g = defaultdict(lambda: [0, 0.0, 0.0])
    for i in range(N):
        if conf(i) and str(ch[i]).strip() == "Гарантия":
            g[str(cl[i]).strip()][0] += 1; g[str(cl[i]).strip()][1] += fnum(tk[i]) or 0; g[str(cl[i]).strip()][2] += fnum(pr[i]) or 0
    loss_n = sum(1 for i in range(N) if conf(i) and str(ch[i]).strip() == "Гарантия" and (fnum(pr[i]) or 0) < 0)
    loss_rub = sum(fnum(pr[i]) for i in range(N) if conf(i) and str(ch[i]).strip() == "Гарантия" and (fnum(pr[i]) or 0) < 0)
    d["guarantee"] = dict(top=sorted(g.items(), key=lambda kv: -kv[1][1])[:10], loss_n=loss_n, loss_rub=loss_rub)

    # продуктивность логистов
    pa = defaultdict(lambda: [0, 0.0]); days = defaultdict(set)
    for i in range(N):
        if not conf(i):
            continue
        a = str(au[i]).strip()
        if not a:
            continue
        pa[a][0] += 1; pa[a][1] += fnum(pr[i]) or 0
        if dl[i]:
            days[a].add(dl[i])
    active = [a for a, v in pa.items() if v[0] >= 20]
    tot_n = sum(pa[a][0] for a in active); tot_d = sum(len(days[a]) for a in active)
    d["logists"] = dict(active=len(active), prod=tot_n / (tot_d or 1),
                        top=sorted(((a, pa[a][0], len(days[a]), pa[a][1]) for a in pa),
                                   key=lambda x: -x[3])[:12])

    # дисбаланс регионов
    lc, uc = Counter(), Counter()
    for i in range(N):
        if not conf(i):
            continue
        if str(rz[i]).strip():
            lc[str(rz[i]).strip()] += 1
        if str(rv[i]).strip():
            uc[str(rv[i]).strip()] += 1
    bal = [(r, lc[r], uc[r], uc[r] - lc[r]) for r in set(lc) | set(uc)]
    d["regions"] = dict(deficit=sorted(bal, key=lambda x: -x[3])[:10],
                        surplus=sorted(bal, key=lambda x: x[3])[:5])
    return d


# ---------- рендер ----------
def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color


def h(doc, text, lvl=1):
    p = doc.add_heading(level=lvl)
    r = p.add_run(text)
    r.font.color.rgb = ACCENT
    r.font.name = "Calibri"
    return p


def para(doc, runs, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, kw in runs:
        set_font(p.add_run(text), **kw)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, hd in enumerate(headers):
        c = t.rows[0].cells[j]
        c.paragraphs[0].clear()
        set_font(c.paragraphs[0].add_run(hd), size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for j, (val, kw) in enumerate(row):
            cells[j].paragraphs[0].clear()
            set_font(cells[j].paragraphs[0].add_run(str(val)), size=10, **kw)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Cm(w)
    return t


def mname(m):
    names = {1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель", 5: "Май"}
    return names.get(m, str(m))


def build(d):
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

    # Титул
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("ВМ-Логистик"), size=14, bold=True, color=GREY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run("Анализ данных по перевозкам"), size=24, bold=True, color=ACCENT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a, b = d["period"]
    set_font(p.add_run(f"Период: {a.strftime('%d.%m.%Y')} — {b.strftime('%d.%m.%Y')}  ·  "
                       f"{d['N']:,} заявок  ·  выгрузка из 1С".replace(",", " ")),
             size=11, italic=True, color=GREY)
    doc.add_paragraph()

    # Резюме
    h(doc, "Краткое резюме", 1)
    fin = d["fin"]
    bullets = [
        (f"Объём: {d['N']:,} заявки за 5 месяцев".replace(",", " "),
         f", {d['status'].get('Подтверждена',0):,} подтверждены ({100*d['status'].get('Подтверждена',0)/d['N']:.0f}%), "
         f"{d['status'].get('Отменена',0)} отменены.".replace(",", " ")),
        ("Выручка ", f"{rub(fin['rev'])} ₽, прибыль {rub(fin['profit'])} ₽, маржа {fin['margin']:.1f}%. "
                     f"Годовой run-rate ≈ 3,4 млрд ₽ выручки."),
        ("Деньги делает СПОТ ", f"— 65% выручки. {d['spot']['rec_n']} повторяющихся маршрутов дают "
                                f"{d['spot']['rec_deals']} заявок и {rub(d['spot']['rec_rev'])} ₽ прибыли — ядро для авто-торгов (контур B)."),
        ("Маржу убивает Гарантия ", f"— 12,7% против 20,5% в среднем. {d['guarantee']['loss_n']} убыточных заявок "
                                    f"на {rub(d['guarantee']['loss_rub'])} ₽; главная дыра — Маревен Фуд (7,7%)."),
        ("Продуктивность логистов ", f"{d['logists']['prod']:.2f} заявки/день — ниже бенчмарка 3,5 и далеко от цели 10 (контур A)."),
        ("Порожний пробег ", "концентрируется на Ростов, СПб, Самару, Краснодар — поле для модуля обратной загрузки."),
    ]
    for head, rest in bullets:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(head), bold=True)
        set_font(p.add_run(rest))

    # 1. Профиль
    h(doc, "1. Профиль данных", 1)
    para(doc, "Один лист, 65 полей. Полнота ключевых полей высокая: тарифы 98%, прибыль 100%, "
              "регионы ~89%, гос-номер 91%. Источник — учётная система 1С.")
    table(doc, ["Показатель", "Значение"], [
        [("Записей всего", {}), (f"{d['N']:,}".replace(",", " "), {"bold": True})],
        [("Уникальных клиентов", {}), (str(d["clients_uni"]), {})],
        [("Уникальных подрядчиков", {}), (f"{d['contractors_uni']:,}".replace(",", " "), {})],
        [("Логистов (авторов)", {}), (str(d["authors_uni"]), {})],
        [("Организации", {}), (" · ".join(f"{k} {v}" for k, v in d["org"].most_common()), {})],
    ], widths=[6, 10])

    # 2. Финансы
    h(doc, "2. Финансовые итоги (подтверждённые заявки)", 1)
    table(doc, ["Метрика", "Значение"], [
        [("Выручка (тариф клиента)", {}), (rub(fin["rev"]) + " ₽", {"bold": True})],
        [("Себестоимость (подрядчик)", {}), (rub(fin["cost"]) + " ₽", {})],
        [("Прибыль", {}), (rub(fin["profit"]) + " ₽", {"bold": True, "color": GOOD})],
        [("Маржа от выручки", {}), (f"{fin['margin']:.1f}%", {"bold": True})],
        [("Средний чек заявки", {}), (rub(fin["avg_check"]) + " ₽", {})],
        [("Средняя прибыль на заявку", {}), (rub(fin["avg_profit"]) + " ₽", {})],
    ], widths=[7, 9])

    para(doc, [("Каналы продаж — где зарабатывается прибыль:", {"bold": True})], space_after=4)
    rows = []
    for k, v in d["channels"]:
        m = 100 * v[2] / v[1] if v[1] else 0
        col = GOOD if m >= 25 else (BAD if m < 13 else None)
        rows.append([(k, {}), (f"{v[0]:,}".replace(",", " "), {}), (rub(v[1]), {}),
                     (rub(v[2]), {}), (f"{m:.1f}%", {"bold": True, "color": col})])
    table(doc, ["Канал", "Заявок", "Выручка ₽", "Прибыль ₽", "Маржа"], rows,
          widths=[3.2, 2.2, 3.6, 3.4, 2.2])
    para(doc, [("СПОТ", {"bold": True}), (" — основной объём (65% выручки), цель контура B. ", {}),
               ("Гарантия", {"bold": True}), (" — самая низкая маржа (12,7%), цель контура A. ", {}),
               ("ПРОЕКТ", {"bold": True}), (" — самые качественные сделки: 155 тыс ₽ прибыли на заявку.", {})], space_after=4)

    # помесячно
    para(doc, [("Помесячная динамика:", {"bold": True})], space_after=4)
    rows = []
    for (y, m), v in d["months"]:
        rows.append([(f"{mname(m)} {y}", {}), (f"{v[0]:,}".replace(",", " "), {}),
                     (rub(v[1]), {}), (rub(v[2]), {}), (f"{100*v[2]/v[1]:.1f}%", {})])
    table(doc, ["Месяц", "Заявок", "Выручка ₽", "Прибыль ₽", "Маржа"], rows,
          widths=[3.2, 2.2, 3.6, 3.4, 2.2])

    # топ-клиенты
    h(doc, "3. Клиенты", 1)
    para(doc, f"{d['clients_uni']} клиентов. Сильная концентрация на FMCG и телеком. Топ-12 по выручке:")
    rows = []
    for c, v in d["top_clients"]:
        m = 100 * v[2] / v[1] if v[1] else 0
        rows.append([(c[:38], {}), (f"{v[0]:,}".replace(",", " "), {}), (rub(v[1]), {}),
                     (f"{m:.1f}%", {"color": BAD if m < 10 else None})])
    table(doc, ["Клиент", "Заявок", "Выручка ₽", "Маржа"], rows, widths=[7.5, 2.0, 3.5, 2.0])

    # 4. СПОТ
    h(doc, "4. СПОТ → контур B (авто-торги)", 1)
    sp = d["spot"]
    para(doc, [(f"{sp['total']:,} СПОТ-заявок раскладываются на {sp['uni']:,} маршрутов. "
                .replace(",", " "), {}),
               (f"Но {sp['rec_n']} маршрутов повторяются ≥10 раз → {sp['rec_deals']} заявок "
                f"(21% потока) и {rub(sp['rec_rev'])} ₽ прибыли", {"bold": True}),
               (" — это ядро, которое имеет смысл автоматизировать через API площадок.", {})])
    rows = [[(r, {}), (f"{v[0]}×", {"bold": True}), (rub(v[1]) + " ₽", {})] for r, v in sp["top"]]
    table(doc, ["Маршрут", "Повторов", "Прибыль"], rows, widths=[8.5, 2.5, 4.0])
    para(doc, [("Нюанс: ", {"bold": True}),
               ("северные/нефтегаз коридоры (Новый Уренгой–Караул, Тюмень–Ноглики) — это проектные "
                "направления под Гарантию/контракт, а не под биржевые торги. Массовые низкомаржинальные "
                "(Наро-Фоминск–Ростов, Орехово-Зуево–СПб) — идеальны под авто-участие в торгах.", {})])

    # 5. Гарантия
    h(doc, "5. Гарантия → контур A (тарифная модель)", 1)
    gr = d["guarantee"]
    para(doc, [("Низкая маржа Гарантии локализована в нескольких клиентах. "
                f"{gr['loss_n']} заявок убыточны, суммарный минус ", {}),
               (rub(gr["loss_rub"]) + " ₽", {"bold": True, "color": BAD}), (".", {})])
    rows = []
    for c, v in gr["top"]:
        m = 100 * v[2] / v[1] if v[1] else 0
        col = BAD if m < 8 else (GOOD if m >= 25 else None)
        rows.append([(c[:34], {}), (f"{v[0]:,}".replace(",", " "), {}), (rub(v[1]), {}),
                     (f"{m:.1f}%", {"bold": True, "color": col})])
    table(doc, ["Клиент (Гарантия)", "Заявок", "Выручка ₽", "Маржа"], rows, widths=[7.5, 2.0, 3.5, 2.0])
    para(doc, [("Маревен Фуд", {"bold": True}),
               (" — главная дыра: 1142 заявки под 7,7%. Нестле (−4,7%) и Мултон (−8,0%) работают в минус. "
                "Это первая точка приложения умного подбора перевозчика: каждый отыгранный процент маржи "
                "по этим клиентам — прямой эффект.", {})])

    # 6. Логисты
    h(doc, "6. Продуктивность логистов → контур A", 1)
    lg = d["logists"]
    para(doc, [(f"{lg['active']} активных логистов (≥20 заявок). Средняя продуктивность ", {}),
               (f"{lg['prod']:.2f} заявки на логиста в день", {"bold": True}),
               (f" — ниже бенчмарка Артёма Герасимова (~3,5) и втрое меньше цели (~10). "
                "Потенциал контура A — ×3,6 по выработке.", {})])
    rows = []
    for a, n, days, p in lg["top"]:
        rows.append([(a, {}), (f"{n:,}".replace(",", " "), {}), (str(days), {}),
                     (f"{n/(days or 1):.1f}", {}), (rub(p) + " ₽", {})])
    table(doc, ["Логист", "Заявок", "Раб. дней", "Заявок/день", "Прибыль"], rows,
          widths=[4.0, 2.0, 2.2, 2.6, 3.5])
    para(doc, [("Важно: ", {"bold": True}),
               ("высокая выработка не равна качеству. Синельников Егор делает 9,8 заявки/день, но его "
                "маржа всего 9,6% — это конвейер по убыточной Гарантии (Маревен). Масштабировать нужно "
                "не объём, а тарифную модель потока.", {})])

    # 7. Обратная загрузка
    h(doc, "7. Обратная загрузка → карта дефицита", 1)
    para(doc, "Регионы, куда много везут, но откуда мало забирают обратный груз (риск порожнего пробега):")
    rows = [[(r, {}), (str(l), {}), (str(u), {}), (f"+{dd}", {"bold": True, "color": BAD})]
            for r, l, u, dd in d["regions"]["deficit"]]
    table(doc, ["Регион", "Загрузка", "Выгрузка", "Дисбаланс"], rows, widths=[6.0, 3.0, 3.0, 3.0])
    para(doc, [("Московская область", {"bold": True}),
               (" — гигантский профицит на входе (отсюда едет всё, обратно мало). Ростов, СПб, Самара, "
                "Краснодар, Татарстан — туда привозят, а обратного груза нет. Именно здесь нужен модуль "
                "поиска обратной загрузки и задача отделу продаж искать грузоотправителей в этих регионах "
                "(идея Артёма Герасимова).", {})])

    # Выводы
    h(doc, "8. Выводы по контурам проекта", 1)
    concl = [
        ("Контур B (авто-торги): ", "бить в СПОТ — 65% выручки. Приоритет — 63 повторяющихся маршрута "
         "(282 млн ₽ прибыли). Низкомаржинальные массовые направления → авто-участие в торгах площадок."),
        ("Контур A (AI-логист): ", "приоритет №1 — тарифная модель Гарантии по Маревену/Нестле/Мултону "
         "(−5,7 млн ₽ на 403 заявках). Второй эффект — рост выработки с 2,76 до целевых ~10 заявок/день."),
        ("Модуль обратной загрузки: ", "карта дефицита по Ростову, СПб, Самаре, Краснодару — порожний "
         "пробег из этих регионов в МО."),
        ("Данные пригодны для дашбордов (контур C): ", "выгрузка 1С содержит всё для командного центра — "
         "финансы, каналы, регионы, логисты, подрядчики."),
    ]
    for head, rest in concl:
        p = doc.add_paragraph(style="List Number")
        set_font(p.add_run(head), bold=True, color=ACCENT)
        set_font(p.add_run(rest))

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18)
    set_font(p.add_run("Источник: выгрузка 1С «данные по перевозкам.xls». "
                       "Расчёты по подтверждённым заявкам. Документ содержит коммерческие данные — "
                       "не передавать без согласования."), size=9, italic=True, color=GREY)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("Сохранено:", OUT)


if __name__ == "__main__":
    build(compute())
