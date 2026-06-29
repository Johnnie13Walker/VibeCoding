#!/usr/bin/env python3
"""Word для собственников: с каких площадок начинать авто-участие в СПОТ-торгах.

Сшивает вкладку «Общая инфо» (площадки на клиента) Нелли с выгрузкой 1С (СПОТ-деньги),
ранжирует площадки по «деньги × реализуемость API», строит план по волнам.
Запуск: .venv/bin/python scripts/build_platforms_docx.py
Выход:  reports/Площадки_авто-торги_СПОТ.docx  (в git не коммитится — коммерческие данные)
"""
import openpyxl, xlrd, re, os
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "reports/Площадки_авто-торги_СПОТ.docx"
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
GOOD = RGBColor(0x1B, 0x7A, 0x3D)
BAD = RGBColor(0xB3, 0x1B, 0x1B)
AMBER = RGBColor(0xB8, 0x7A, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)


def keyize(s):
    s = str(s).lower().replace("ё", "е")
    s = re.sub(r"и|й", "и", s)
    s = re.sub(r"э|е", "е", s)
    for w in ["ооо", "оао", "зао", "пао", "ао", "холдингс", "россия", "партнерс",
              "сводныи", "сэнтрал", "фуд", "компания", "пивоваренная", "снабжение"]:
        s = s.replace(w, "")
    return re.sub(r"[^a-zа-я0-9]", "", s)


def platform(op, url):
    s = f"{op} {url}".lower()
    if "atrucks" in s or "атракс" in s:
        return "Атракс"
    if ("logist" in s and "pro" in s) or "логистпро" in s or "логист про" in s or "logistpro" in s:
        return "Логист-Про"
    if "torgtrans" in s or "торгтранс" in s:
        return "ТоргТранс"
    if "loginet" in s or "логинет" in s:
        return "Логинет"
    if "ati" in s and ".su" in s:
        return "ATI.SU"
    if "tms" in s or "trucker" in s or "труке" in s or "сберкорус" in s:
        return "TMS/Trucker"
    if "bidzaar" in s or "бидзар" in s:
        return "Bidzaar"
    if "почт" in s or "@" in url or url.strip() == "-" or op.strip().lower() == "тг":
        return "Почта/мессенджер"
    return "Свой портал"


API = {
    "Атракс": "✅ из коробки", "Логист-Про": "✅ из коробки", "ATI.SU": "✅ зрелый",
    "ТоргТранс": "по договору", "Логинет": "по договору", "Bidzaar": "под заказчика",
    "TMS/Trucker": "клиентский", "Почта/мессенджер": "нет (парсинг)", "Свой портал": "только RPA",
}


def compute():
    wb = openpyxl.load_workbook("data/clients_nelli.xlsx", data_only=True)
    ws = wb["Общая инфо "]
    rows = []
    for r in range(2, ws.max_row + 1):
        nm = ws.cell(r, 1).value
        if nm and str(nm).strip():
            rows.append((str(nm).strip(), str(ws.cell(r, 15).value or ""),
                         str(ws.cell(r, 16).value or ""), str(ws.cell(r, 17).value or "")))

    sh = xlrd.open_workbook("data/perevozki_raw.xls").sheet_by_index(0)
    H = {str(sh.cell_value(0, c)).strip(): c for c in range(sh.ncols)}

    def fnum(x):
        s = str(x).strip().replace("\xa0", "").replace(" ", "").replace(",", ".")
        try:
            return float(s)
        except Exception:
            return None

    spot = defaultdict(lambda: [0, 0.0])
    for r in range(1, sh.nrows):
        if str(sh.cell_value(r, H["Статус заявки"])).strip() != "Подтверждена":
            continue
        if str(sh.cell_value(r, H["Канал продаж"])).strip() != "СПОТ":
            continue
        c = str(sh.cell_value(r, H["Клиент"])).strip()
        spot[c][0] += 1
        spot[c][1] += fnum(sh.cell_value(r, H["Прибыль по клиенту руб"])) or 0
    full_keys = [(keyize(c), c) for c in spot]

    def money(short):
        k = keyize(short)
        if not k:
            return 0, 0.0
        for fk, full in full_keys:
            if k in fk or fk in k or (len(k) >= 4 and k[:5] in fk):
                return spot[full]
        return 0, 0.0

    agg = defaultdict(lambda: [set(), 0, 0.0])
    HOT = re.compile(r"актив|переб|по ставк|не проход|необходим", re.I)
    hot = defaultdict(int)
    for nm, op, url, stt in rows:
        p = platform(op, url)
        d, pr = money(nm)
        agg[p][0].add(nm)
        agg[p][1] += d
        agg[p][2] += pr
        if HOT.search(stt):
            hot[p] += 1

    total = sum(v[1] for v in spot.values())
    return dict(platforms=sorted(agg.items(), key=lambda kv: -kv[1][2]), hot=hot,
                total_spot=total, spot=spot, money=money)


def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = color


def h(doc, text, lvl=1):
    p = doc.add_heading(level=lvl)
    rr = p.add_run(text); rr.font.color.rgb = ACCENT; rr.font.name = "Calibri"
    return p


def para(doc, runs, space_after=6, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    if isinstance(runs, str):
        runs = [(runs, {})]
    for t, kw in runs:
        set_font(p.add_run(t), **kw)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, hd in enumerate(headers):
        c = t.rows[0].cells[j]; c.paragraphs[0].clear()
        set_font(c.paragraphs[0].add_run(hd), size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for j, (val, kw) in enumerate(row):
            cells[j].paragraphs[0].clear()
            set_font(cells[j].paragraphs[0].add_run(str(val)), size=10, **kw)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Cm(w)
    return t


def rub(x):
    return f"{x:,.0f}".replace(",", " ")


def build(d):
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("ВМ-Логистик · проект внедрения ИИ"), size=13, bold=True, color=GREY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run("Авто-участие в СПОТ-торгах: с каких площадок начинать"), size=22, bold=True, color=ACCENT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Приоритизация по принципу «деньги × реализуемость». "
                       "Основано на выгрузке 1С (янв–май 2026) и реестре площадок КМ."),
             size=10, italic=True, color=GREY)
    doc.add_paragraph()

    # Резюме
    h(doc, "Главное", 1)
    for head, rest in [
        ("СПОТ — это 65% выручки.", f" Прибыль по СПОТ {rub(d['total_spot'])} ₽ за 5 мес (≈{rub(d['total_spot']*2.4)} ₽/год). "
                                    "Именно сюда бьёт авто-участие в торгах."),
        ("Главный потенциал — упускаемый рынок.", " По оценке собственников вручную вы участвуете лишь в ~6% спот-заявок рынка "
                                                  "и выигрываете ~22% из них. 94% потока проходит мимо — здесь и лежит цель ×3."),
        ("Эффект ≠ просто «где больше денег».", " Самые жирные площадки технически самые тяжёлые (свои порталы — только RPA, "
                                                "ТоргТранс — API по договору). Начинать нужно там, где API готов из коробки."),
        ("Рекомендация:", " старт — Атракс + Логист-Про (API из коробки), стратегический рост — ATI.SU (открытый рынок), "
                          "вторая волна за деньгами — ТоргТранс."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(head), bold=True)
        set_font(p.add_run(rest))

    # Карта площадок
    h(doc, "1. Где крутится СПОТ-прибыль по площадкам", 1)
    para(doc, "Каждый клиент размещает заявки на своей площадке. Сшивка реестра КМ с деньгами из 1С "
              "(суммы приблизительные — клиенты сопоставлены по совпадению наименований):")
    rows = []
    for pname, v in d["platforms"]:
        cnt = len(v[0]); pr = v[2]
        apit = API.get(pname, "?")
        col = GOOD if "✅" in apit else (AMBER if apit in ("по договору", "под заказчика", "клиентский") else BAD)
        rows.append([(pname, {}), (str(cnt), {}), (rub(pr), {"bold": pr > 15e6}),
                     (rub(pr * 2.4), {}), (apit, {"color": col, "bold": "✅" in apit})])
    table(doc, ["Площадка", "Клиентов", "СПОТ-приб/5мес ₽", "≈/год ₽", "API"], rows,
          widths=[3.6, 2.0, 3.6, 3.2, 3.2])
    para(doc, [("Чтение таблицы: ", {"bold": True}),
               ("«Свой портал» и «ТоргТранс» держат больше всего денег, но это десятки разных порталов "
                "(RPA) либо доступ по договору с площадкой. «Атракс» — один коннектор закрывает 15 клиентов "
                "сразу, API из коробки.", {})], space_after=4)

    # План волн
    h(doc, "2. План по волнам", 1)
    waves = [
        ("Волна 1 — Атракс + Логист-Про", GOOD,
         "API из коробки, согласований с площадкой не требуется. Атракс = один интегратор на 15 клиентов — "
         "лучший охват на единицу разработки и идеальный полигон для обучения модели скоринга и авто-ставки. "
         "Логист-Про добавляет Мултон (в Гарантии он у нас убыточен — авто-торг на СПОТ даёт двойной эффект). "
         "Срок до первого результата — недели."),
        ("Волна 1-параллельно — ATI.SU", GOOD,
         "Зрелый API. В реестре 1 клиент, но ATI — это биржа открытого рынка, а не закреплённый клиент. "
         "Именно здесь лежит потенциал ×3: выход в 94% потока заявок, который сейчас проходит мимо."),
        ("Волна 2 — ТоргТранс", AMBER,
         "Самые концентрированные деньги (ИнБев + Маревен — почти вся сумма площадки), но API только по "
         "договору с площадкой. Параллельно с Волной 1 запускаем переговоры о доступе; как только дают — "
         "один интегратор закрывает почти всю сумму."),
        ("Точечно — Стройпроектсервис", AMBER,
         "Один клиент даёт 30 млн ₽ = 17% всей СПОТ-прибыли, работает на своей площадке и отсутствует в "
         "реестре КМ. Это не «площадка», а отдельный RPA-робот под одного крупного — окупается сам."),
    ]
    for i, (title, col, body) in enumerate(waves, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(f"{i}. {title}"), size=12, bold=True, color=col)
        para(doc, body, space_after=8)

    # Экономика Волны 1
    h(doc, "3. Экономика Волны 1 (расчёт по маршрутам)", 1)
    w = d["wave1"]
    table(doc, ["Показатель", "Значение"], [
        [("СПОТ-заявок клиентов Волны 1 (5 мес)", {}), (f"{w['deals']:,}".replace(",", " "), {})],
        [("СПОТ-прибыль Волны 1", {}), (rub(w["profit"]) + " ₽ (≈" + rub(w["profit"] * 2.4) + " ₽/год)", {"bold": True})],
        [("в т.ч. Атракс", {}), (rub(w["atraks"]) + " ₽", {})],
        [("в т.ч. Логист-Про", {}), (rub(w["lp"]) + " ₽", {})],
        [("Повторяющихся маршрутов (≥5 раз)", {}), (f"{w['rec_routes']} маршрутов = {w['rec_deals']} заявок ({w['rec_share']:.0f}% потока)", {})],
    ], widths=[8.0, 8.0])
    para(doc, [("Прямой эффект (нижняя граница): ", {"bold": True}),
               (f"авто-участие добавляет +20–30% объёма на повторяемых маршрутах = "
                f"+{rub(w['uplift20']*2.4)}…{rub(w['uplift30']*2.4)} ₽/год. ", {}),
               ("Это только расширение уже выигрываемых маршрутов.", {"italic": True})], space_after=4)
    para(doc, [("Главный эффект (выше прямого): ", {"bold": True}),
               ("(1) рост доли побед в аукционах, которые логист сейчас физически не успевает отрабатывать; "
                "(2) высвобождение времени логистов (сейчас 2,76 заявки/день при цели 10) под более "
                "маржинальные сделки. Эти два рычага кратно больше расширения существующих маршрутов.", {})])
    para(doc, [("Топ повторяемых маршрутов Волны 1 (полигон для старта): ", {"bold": True})], space_after=4)
    rows = [[(r, {}), (f"{n}×", {"bold": True}), (rub(pr) + " ₽", {"color": BAD if pr < 0 else None})]
            for r, n, pr in w["top_routes"]]
    table(doc, ["Маршрут", "Повторов/5мес", "Прибыль"], rows, widths=[8.5, 3.0, 3.5])

    # Риски/запросы
    h(doc, "4. Что нужно для старта", 1)
    for t in [
        "Dev-доступ к API Атракс и Логист-Про от лица перевозчика (логины уже есть — нужен технический доступ).",
        "Уточнить у КМ площадки крупных СПОТ-клиентов вне реестра: Стройпроектсервис (30 млн), Иркутский завод полимеров, Газпромнефть-Снабжение, НПО Утёсов.",
        "Правило флор-цены (минимальная маржа) для авто-ставки — чтобы робот не выигрывал убыточные рейсы (на маршруте Щёлково–Оренбург уже видим минус).",
        "Политика: авто-торги — это не контроль логиста (на чём настаивает Голубев), а отдельный модуль участия в аукционах. Развести явно.",
    ]:
        para(doc, t, style="List Bullet", space_after=4)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16)
    set_font(p.add_run("Источники: выгрузка 1С «данные по перевозкам» (янв–май 2026, подтверждённые заявки), "
                       "реестр площадок КМ (вкладка «Общая инфо»). Суммы по площадкам приблизительны "
                       "(сопоставление клиентов по наименованию). Документ содержит коммерческие данные."),
             size=9, italic=True, color=GREY)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("Сохранено:", OUT)


def wave1(d):
    """Экономика Атракс + Логист-Про по маршрутам."""
    wb = openpyxl.load_workbook("data/clients_nelli.xlsx", data_only=True)
    ws = wb["Общая инфо "]
    keys = []
    for r in range(2, ws.max_row + 1):
        nm = ws.cell(r, 1).value
        if not (nm and str(nm).strip()):
            continue
        p = platform(str(ws.cell(r, 15).value or ""), str(ws.cell(r, 16).value or ""))
        if p in ("Атракс", "Логист-Про"):
            keys.append((keyize(nm), p))
    sh = xlrd.open_workbook("data/perevozki_raw.xls").sheet_by_index(0)
    H = {str(sh.cell_value(0, c)).strip(): c for c in range(sh.ncols)}

    def fnum(x):
        s = str(x).strip().replace("\xa0", "").replace(" ", "").replace(",", ".")
        try:
            return float(s)
        except Exception:
            return None

    def mp(client):
        k = keyize(client)
        for ck, p in keys:
            if ck and (ck in k or k in ck or (len(ck) >= 4 and ck[:5] in k)):
                return p
        return None

    deals = 0; profit = 0.0; atraks = 0.0; lp = 0.0
    routes = defaultdict(lambda: [0, 0.0])
    for r in range(1, sh.nrows):
        if str(sh.cell_value(r, H["Статус заявки"])).strip() != "Подтверждена":
            continue
        if str(sh.cell_value(r, H["Канал продаж"])).strip() != "СПОТ":
            continue
        p = mp(str(sh.cell_value(r, H["Клиент"])).strip())
        if not p:
            continue
        pr = fnum(sh.cell_value(r, H["Прибыль по клиенту руб"])) or 0
        rt = str(sh.cell_value(r, H["Маршрут"])).strip()
        deals += 1; profit += pr
        atraks += pr if p == "Атракс" else 0
        lp += pr if p == "Логист-Про" else 0
        if rt:
            routes[rt][0] += 1; routes[rt][1] += pr
    rec = [(r, v[0], v[1]) for r, v in routes.items() if v[0] >= 5]
    rec_deals = sum(x[1] for x in rec)
    ap = sum(x[2] for x in rec) / rec_deals if rec_deals else 0
    return dict(deals=deals, profit=profit, atraks=atraks, lp=lp,
                rec_routes=len(rec), rec_deals=rec_deals, rec_share=100 * rec_deals / deals,
                uplift20=rec_deals * 0.20 * ap, uplift30=rec_deals * 0.30 * ap,
                top_routes=sorted(rec, key=lambda x: -x[1])[:10])


if __name__ == "__main__":
    d = compute()
    d["wave1"] = wave1(d)
    build(d)
