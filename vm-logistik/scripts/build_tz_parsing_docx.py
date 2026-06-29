#!/usr/bin/env python3
"""ТЗ для подрядчиков (AI-разработчики): Этап 1 — парсинг заявок в ЛК логиста.

Документ для встречи Жени с партнёрами. Цель Этапа 1 — увидеть объём заявок,
которые мы не забираем, через парсинг опишек/агрегаторов в личный кабинет логиста.
Цифры объёма подтягиваются из вкладки «СПОТы» (реальные данные торгов).

Запуск: .venv/bin/python scripts/build_tz_parsing_docx.py
Выход:  reports/ТЗ_Этап1_парсинг_заявок_ЛК-логиста.docx (в git не коммитится)
"""
import openpyxl, re, os
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "reports/ТЗ_Этап1_парсинг_заявок_ЛК-логиста.docx"
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
GOOD = RGBColor(0x1B, 0x7A, 0x3D)
BAD = RGBColor(0xB3, 0x1B, 0x1B)
GREY = RGBColor(0x60, 0x60, 0x60)


def keyize(s):
    s = str(s).lower().replace("ё", "е"); s = re.sub(r"и|й", "и", s); s = re.sub(r"э|е", "е", s)
    return re.sub(r"[^a-zа-я0-9]", "", s)


def platform(op, url):
    s = f"{op} {url}".lower()
    if "atrucks" in s or "атракс" in s: return "Атракс"
    if ("logist" in s and "pro" in s) or "логистпро" in s or "логист про" in s or "logistpro" in s: return "Логист-Про"
    return None


def spot_volume():
    wb = openpyxl.load_workbook("data/clients_nelli.xlsx", data_only=True)
    ws = wb["СПОТы"]
    def num(v):
        try:
            return float(v)
        except Exception:
            return 0.0
    spotv = defaultdict(lambda: [0.0, 0.0, 0.0]); tot = [0.0, 0.0, 0.0]; months = set()
    import datetime
    for r in range(4, ws.max_row + 1):
        c = ws.cell(r, 5).value; f = ws.cell(r, 6).value
        if not c or f is None: continue
        z = num(f); u = num(ws.cell(r, 7).value); w = num(ws.cell(r, 8).value)
        if z == 0 and u == 0 and w == 0: continue
        spotv[str(c).strip()][0] += z; spotv[str(c).strip()][1] += u; spotv[str(c).strip()][2] += w
        tot[0] += z; tot[1] += u; tot[2] += w
        d = ws.cell(r, 4).value
        if isinstance(d, datetime.datetime): months.add((d.year, d.month))
    M = len(months) or 9
    ws2 = wb["Общая инфо "]
    reg = {"Атракс": [], "Логист-Про": []}
    for r in range(2, ws2.max_row + 1):
        nm = ws2.cell(r, 1).value
        if not (nm and str(nm).strip()): continue
        p = platform(str(ws2.cell(r, 15).value or ""), str(ws2.cell(r, 16).value or ""))
        if p: reg[p].append(str(nm).strip())
    keys = [(keyize(c), c) for c in spotv]
    def lk(name):
        k = keyize(name)
        for sk, sc in keys:
            if sk and (sk in k or k in sk or (len(sk) >= 3 and len(k) >= 4 and sk[:4] == k[:4])): return sc
        return None
    res = {}
    for P in ("Атракс", "Логист-Про"):
        seen = set(); z = u = w = 0.0
        for nm in set(reg[P]):
            sc = lk(nm)
            if sc and sc not in seen:
                seen.add(sc); z += spotv[sc][0]; u += spotv[sc][1]; w += spotv[sc][2]
        res[P] = dict(z_pm=z / M, part=100 * u / z if z else 0, win=100 * w / z if z else 0, notclosed_pm=(z - w) / M)
    return dict(months=M, tot_z=tot[0], part=100 * tot[1] / tot[0], win=100 * tot[2] / tot[0], plat=res)


def sf(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic; run.font.name = "Calibri"
    if color is not None: run.font.color.rgb = color


def h(doc, t, lvl=1):
    p = doc.add_heading(level=lvl); r = p.add_run(t); r.font.color.rgb = ACCENT; r.font.name = "Calibri"; return p


def para(doc, runs, sa=6, style=None):
    p = doc.add_paragraph(style=style); p.paragraph_format.space_after = Pt(sa)
    if isinstance(runs, str): runs = [(runs, {})]
    for t, kw in runs: sf(p.add_run(t), **kw)
    return p


def bullet(doc, runs, sa=3):
    return para(doc, runs, sa=sa, style="List Bullet")


def num_item(doc, runs, sa=3):
    return para(doc, runs, sa=sa, style="List Number")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, hd in enumerate(headers):
        c = t.rows[0].cells[j]; c.paragraphs[0].clear()
        sf(c.paragraphs[0].add_run(hd), size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for j, (val, kw) in enumerate(row):
            cells[j].paragraphs[0].clear(); sf(cells[j].paragraphs[0].add_run(str(val)), size=10, **kw)
    if widths:
        for j, wd in enumerate(widths):
            for r in t.rows: r.cells[j].width = Cm(wd)
    return t


def build(sv):
    doc = Document()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Cm(1.7); sec.bottom_margin = Cm(1.7); sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run("ВМ-Логистик · проект внедрения ИИ"), size=13, bold=True, color=GREY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    sf(p.add_run("Техническое задание · Этап 1\nПарсинг заявок в личный кабинет логиста"), size=21, bold=True, color=ACCENT)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run("Для команды AI-разработчиков. Цель этапа — увидеть объём заявок, которые компания "
                 "сейчас не забирает."), size=10, italic=True, color=GREY)
    doc.add_paragraph()

    # 1. Контекст и цель
    h(doc, "1. Контекст и цель", 1)
    para(doc, [("СПОТ-торги дают основную маржу компании, но обрабатываются вручную. По данным учёта за "
                f"{sv['months']} месяцев через площадки прошло ", {}),
               (f"{sv['tot_z']:,.0f} заявок".replace(",", " "), {"bold": True}),
               (f", из них мы участвуем в {sv['part']:.1f}% и выигрываем {sv['win']:.1f}%. ", {}),
               ("Остальное проходит мимо — логист физически не успевает даже посмотреть поток.", {})])
    para(doc, [("Корень проблемы: ", {"bold": True}),
               ("листинг заявок (первичный слой данных) сейчас не виден нигде в едином месте. В 1С попадает "
                "лишь ~3% (по отменённым заявкам). Пока мы не видим полный поток, мы не можем ни оценить "
                "упущенную выгоду, ни управлять выбором заявок.", {})])
    para(doc, [("Цель Этапа 1: ", {"bold": True, "color": ACCENT}),
               ("собрать (спарсить) заявки с приоритетных площадок в единый листинг внутри личного кабинета "
                "логиста (ЛК) и показать объём. Это «низковисящие яблоки» — технически самый быстрый шаг. "
                "Оценка ликвидности, авто-ставка и голосовой ассистент — следующие этапы, в это ТЗ не входят.", {})])

    # 2. Целевая картина — ЛК логиста
    h(doc, "2. Целевая картина: личный кабинет логиста (ЛК)", 1)
    para(doc, "ЛК логиста — единая система-контейнер, в которую поэтапно заводится вся работа со СПОТ-заявками. "
              "Полное видение (для понимания направления; в Этап 1 входит только пункт 1):")
    for i, t in enumerate([
        ("Единый листинг заявок. ", "Логист видит заявки со всех подключённых площадок в одном окне: направление, "
         "дата отгрузки, тоннаж, тип ТС, ставка, клиент, дедлайн. Фильтры по направлению/тоннажу/клиенту. "
         "(← Этап 1)"),
        ("Оценка ликвидности. ", "Каждая заявка помечается «рентабельно / нерентабельно» на основе исторических "
         "данных и тарифного коридора. (Этап 2)"),
        ("Публикация и апдейт груза. ", "Кнопка «опубликовать» по заявке с автозаполнением параметров; апдейт "
         "размещённых грузов в АТИ (чтобы поднимались выше) — внутри ЛК, а не вне его. Цель — чтобы логист "
         "работал только отсюда. (Этап 3)"),
        ("Снижение зависимости от логиста. ", "Голосовой/авто-ассистент берёт заявки, которые логист не успевает "
         "обработать. (Этап 4)"),
    ], 1):
        p = doc.add_paragraph(style="List Number"); sf(p.add_run(t[0]), bold=True); sf(p.add_run(t[1]))

    # 3. Объём Этапа 1
    h(doc, "3. Что парсим в Этапе 1 (приоритет)", 1)
    para(doc, "Начинаем с площадок с открытым доступом / API («опишки») и агрегаторов — их проще всего "
              "реализовать технически. Ниже — реальный объём по двум стартовым площадкам:")
    rows = []
    for P in ("Атракс", "Логист-Про"):
        pl = sv["plat"][P]
        rows.append([(P, {"bold": True}), (f"{pl['z_pm']:,.0f}".replace(",", " "), {}),
                     (f"{pl['part']:.1f}%", {}), (f"{pl['win']:.1f}%", {}),
                     (f"{pl['notclosed_pm']:,.0f}".replace(",", " "), {"bold": True, "color": BAD})])
    table(doc, ["Площадка", "Заявок/мес", "Участвуем", "Выигрываем", "Не закрываем/мес"], rows,
          widths=[4.0, 3.0, 3.0, 3.0, 3.5])
    para(doc, [("Приоритет источников: ", {"bold": True}),
               ("(1) Атракс — API из коробки; (2) Логист-Про — открытая опишка, забирает заявки многих клиентов; "
                "(3) агрегаторы заявок (уточнить у Нелли/Артёма конкретные сервисы). Свои порталы крупных клиентов "
                "(зоопарк, у каждого своё) — позже и по одному, начиная с самых объёмных.", {})])
    para(doc, [("Поля заявки для парсинга (минимум): ", {"bold": True}),
               ("клиент/грузовладелец, направление (город загрузки → город выгрузки), дата отгрузки, тоннаж, "
                "тип ТС, ставка/тариф, дедлайн подачи, ссылка на лот, площадка-источник, время появления.", {})])

    # 4. Результат Этапа 1
    h(doc, "4. Ожидаемый результат Этапа 1", 1)
    for t in [
        "Работающий парсер заявок с Атракс и Логист-Про (+ хотя бы один агрегатор), запускаемый по расписанию.",
        "Блок «Листинг заявок» в ЛК логиста: таблица заявок с фильтрами, обновляется автоматически.",
        "Видимость объёма: сколько заявок/мес проходит, по направлениям/клиентам — то, что сейчас не видно нигде.",
        "Документированная оценка по каждой площадке: доступ через API или парсинг, ограничения, трудозатраты.",
    ]:
        num_item(doc, t)
    para(doc, [("Метрика успеха: ", {"bold": True, "color": GOOD}),
               ("на выходе видим полный поток заявок ключевых площадок и можем сесть с Артёмом Герасимовым, "
                "чтобы оценить «куски пирога, которые не забираем».", {})], sa=4)

    # 5. Технические вопросы и риски
    h(doc, "5. Что нужно проработать подрядчику", 1)
    for t in [
        ("Способ доступа по каждой площадке. ", "Изучить API-документацию Атракс / Логист-Про / агрегаторов. Где "
         "API удобный — работаем через него; где кривой или закрытый — парсинг. Определить в ходе реализации."),
        ("Доступы. ", "Какие учётные данные/токены нужны от ВМ-Логистик (логины перевозчика уже есть — нужен "
         "технический/dev-доступ)."),
        ("Интерактивность. ", "На Этапе 1 — только чтение (получить и показать заявки). Возможность «забрать» "
         "заявку обратно на площадку — отдельный вопрос следующего этапа."),
        ("Оценка работ. ", "Трудозатраты, срок и бюджет по Этапу 1. Ожидание: недорого, т.к. это интеграция с "
         "опишками, а не сложный AI."),
    ]:
        p = doc.add_paragraph(style="List Bullet"); sf(p.add_run(t[0]), bold=True); sf(p.add_run(t[1]))

    # 6. Чего НЕ делаем сейчас
    h(doc, "6. Вне рамок Этапа 1 (важно)", 1)
    for t in [
        "Голосовой ассистент / исходящие звонки — отложено (нет надёжного канала коммуникации в РФ; преждевременно).",
        "Развитие коммуникации с логистами/подрядчиками через мессенджеры — отложено по той же причине.",
        "Авто-ставка и автоматическое участие в торгах — следующие этапы, после того как увидим и оценим поток.",
        "Подключение своих порталов крупных клиентов — позже, по одному, начиная с самых объёмных.",
    ]:
        bullet(doc, t)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16)
    sf(p.add_run("Решение принято на встрече с Романом (учредитель, направление AI-логиста). Последовательность: "
                 "парсинг → оценка ликвидности → завод логистов в ЛК → снижение зависимости от логиста. "
                 "Объёмы — по данным учёта КМ (вкладка «СПОТы»)."), size=9, italic=True, color=GREY)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("Сохранено:", OUT)


if __name__ == "__main__":
    build(spot_volume())
